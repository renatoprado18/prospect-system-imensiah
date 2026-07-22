"""
RACI Smart Updates — interpretacao livre de mensagens do grupo WhatsApp.

Substitui o filtro de keyword + regex puro do parse_raci_update por classificador
Claude que aceita texto livre (e futuramente docs). Mantem parse_raci_update
como fast-path: regex roda primeiro, AI so dispara se regex nao matchar.

Fluxo:
1. Webhook recebe msg do grupo Conselho
2. Tenta parse_raci_update (regex) — se match, aplica
3. Se nao match, chama propose_updates_from_text (AI)
4. Filtra por confianca: alta = auto-apply; media = notif Renato; baixa = ignora
5. Aplica alta confianca + responde no grupo
6. Notifica Renato no WA privado se houver propostas de media confianca

Cost: ~Haiku 4.5 ~$0.001/msg. Grupo ativo 20 msg/dia = ~$0.6/mes/grupo.
"""
from __future__ import annotations

import os
from services import llm
from services import llm_usage
import json
import logging
import re
from datetime import datetime
from typing import Dict, List, Optional

import httpx

logger = logging.getLogger(__name__)

CONSELHOOS_DATABASE_URL = os.getenv("CONSELHOOS_DATABASE_URL", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# Tamanho minimo de msg pra rodar AI (filtra "ok", "👍", reactions)
MIN_TEXT_LEN_FOR_AI = 12

# Phase 2: limite anti-loop pra docs gigantes baixados da Evolution.
MAX_MEDIA_BYTES = 20 * 1024 * 1024  # 20MB

# ── Guardrails semanticos (fix 22/07 — FP de cortesia fechando item de revisao) ──
# BUG observado ao vivo: mensagens de CORTESIA pos-reuniao ("fico feliz",
# "grato pela colaboracao", "agradeço os feedbacks da reuniao") fecharam o item
# "Revisar o questionario dos socios" — nenhuma falava do documento em si.
# Raiz: o classifier tratava atividade generica/social como conclusao, sem checar
# se a mensagem satisfaz ESTE entregavel especifico. Estas camadas deterministicas
# rodam DEPOIS do LLM como rede de seguranca testavel.
#
# KILL-SWITCH: por padrao, itens de JULGAMENTO (revisao/aprovacao/avaliacao) NUNCA
# auto-fecham (viram no maximo proposta 'media' = human-in-loop). Pra reabilitar
# auto-close desses itens (nao recomendado), set RACI_JUDGMENT_AUTOCLOSE=1.


def _judgment_autoclose_enabled() -> bool:
    """Kill-switch. Default OFF: item de julgamento nunca auto-fecha (so propoe)."""
    return (os.getenv("RACI_JUDGMENT_AUTOCLOSE", "0") or "").strip().lower() in ("1", "true", "yes", "on")


def _strip_accents(s: str) -> str:
    import unicodedata
    return "".join(c for c in unicodedata.normalize("NFD", s or "") if unicodedata.category(c) != "Mn")


def _norm(s: str) -> str:
    """lowercase + sem acento + sem prefixo [SENDER em ...] do batch."""
    s = s or ""
    # Remove prefixo de contexto que o batch injeta: "[Nome em 12/07 14:30]\n..."
    s = re.sub(r"^\s*\[[^\]]{0,60}\]\s*", "", s, count=1)
    return _strip_accents(s).lower().strip()


# Marcadores de cortesia/gratidao pos-reuniao — sinal INSUFICIENTE pra fechar item.
_COURTESY_MARKERS = [
    "fico feliz", "fico muito feliz", "muito feliz", "que feliz",
    "feliz em colaborar", "feliz em ajudar", "feliz em participar",
    "obrigad", "muito grat", "grato ", "grata ", "agradec", "agradeco",
    "parabens", "que bom", "maravilh", "excelente reuniao", "otima reuniao",
    "otima conversa", "otimo encontro", "foi um prazer", "prazer em",
    "um abraco", "abracos", "abraco a todos", "adorei", "amei",
    "conte comigo", "contem comigo", "a disposicao", "estou a disposicao",
    "estamos juntos", "sucesso a todos", "bom te ver", "bom ter",
    "teras cuidado", "todo o cuidado", "todo cuidado", "cuidado com",
    "colaborar com voce", "colaborar com voces", "grande abraco",
]

# Sinal SUBSTANTIVO — indica que a msg fala de trabalho/entregavel concreto,
# nao so social. Se presente, a msg NAO e cortesia-pura.
_SUBSTANTIVE_MARKERS = [
    "document", "questionario", "planilha", "relatorio", "anexo", "versao",
    "revisei", "revisado", "revisamos", "aprovado", "aprovei", "assinado",
    "enviei", "enviado", "envio o", "envio a", "segue o", "segue a", "segue em",
    "minuta", "contrato", "proposta", "ata ", "prazo", "deadline", "reais",
    "r$", "kommo", "crm", "planejamento", "orcamento", "cronograma", "draft",
    "preenchi", "preenchido", "respondi o", "respondemos", "finalizei",
    "conclui ", "concluido o", "concluida a", "esta pronto", "esta pronta",
    "ficou pronto", "atualizei", "subi o", "subi a", "compartilhei",
]


def _is_courtesy_only(text: str) -> bool:
    """True quando a msg e essencialmente cortesia/gratidao pos-reuniao SEM
    referencia a nenhum entregavel/trabalho concreto. Estas nao devem mover
    NENHUM item de RACI (nem abrir proposta).

    Ex FP (22/07): 'fico feliz em colaborar', 'agradeco os feedbacks da reuniao',
    'fico feliz, teras todo o cuidado' -> True (cortesia).
    Ex legitima: 'obrigado, ja revisei o questionario e esta aprovado' -> False
    (tem 'revisei'/'questionario'/'aprovado' = substantivo)."""
    n = _norm(text)
    if not n:
        return False
    has_courtesy = any(m in n for m in _COURTESY_MARKERS)
    if not has_courtesy:
        return False
    has_substantive = any(m in n for m in _SUBSTANTIVE_MARKERS)
    return not has_substantive


# Itens de JULGAMENTO: revisar/aprovar/avaliar/validar/analisar/deliberar.
# Nao fecham por "atividade" — so por evidencia de que O ENTREGAVEL foi julgado.
_JUDGMENT_MARKERS = [
    "revisar", "revisao", "revise ", "revisao do", "revisao da",
    "aprovar", "aprovacao", "aprovaco", "aprove ", "homologar",
    "validar", "validacao", "valide ", "avaliar", "avaliacao",
    "analisar", "analise", "parecer", "deliberar", "deliberacao",
    "julgar", "opinar", "opiniao sobre", "decidir sobre", "definir sobre",
    "considerar a", "considerar o", "dar retorno sobre", "dar feedback sobre",
]

_STOPWORDS = {
    "o", "a", "os", "as", "de", "do", "da", "dos", "das", "e", "em", "no", "na",
    "nos", "nas", "um", "uma", "para", "pra", "por", "com", "sem", "sobre", "ao",
    "aos", "que", "se", "sua", "seu", "suas", "seus", "the", "of", "to", "com",
}


def _is_judgment_item(acao: str) -> bool:
    """True se o item RACI e de revisao/aprovacao/julgamento."""
    n = _norm(acao)
    return any(m in n for m in _JUDGMENT_MARKERS)


def _item_keywords(acao: str) -> set:
    """Substantivos salientes do item (>=4 chars, sem stopword) pra checar match."""
    n = _norm(acao)
    toks = re.findall(r"[a-z0-9]{4,}", n)
    # Descarta os proprios verbos de julgamento pra nao dar match trivial
    verbs = {"revisar", "revisao", "aprovar", "aprovacao", "validar", "avaliar",
             "avaliacao", "analisar", "analise", "parecer", "deliberar", "julgar",
             "conselho", "diretora", "executiva", "socios", "socio"}
    return {t for t in toks if t not in _STOPWORDS and t not in verbs}


def _references_deliverable(text: str, evidencia: str, acao: str) -> bool:
    """True se a msg (ou a evidencia extraida) menciona o entregavel nomeado no
    item. Ex: item 'Revisar o questionario dos socios' + msg com 'questionario'
    -> True. 'fico feliz em colaborar' -> False."""
    kws = _item_keywords(acao)
    if not kws:
        return False
    hay = _norm((text or "") + " " + (evidencia or ""))
    return any(k in hay for k in kws)


def _apply_guardrails(proposals: List[Dict], items_by_id: Dict[str, str], text: str) -> List[Dict]:
    """Rede de seguranca deterministica pos-LLM. items_by_id: item_id -> acao.

    Regras (anti-gen-1 / human-in-loop):
    1. Msg de cortesia-pura -> descarta TODAS as propostas (sinal insuficiente).
    2. Item de julgamento (revisao/aprovacao) + conclusao:
       - sem referencia ao entregavel -> DESCARTA (nao fecha revisao por atividade).
       - com referencia -> rebaixa confianca pra 'media' (PROPOR, nao executa),
         a menos que RACI_JUDGMENT_AUTOCLOSE=1.
    3. Itens operacionais com evidencia clara: inalterados (auto-close preservado).
    """
    if _is_courtesy_only(text):
        logger.info("smart_updates: msg cortesia-pura, descartando %d proposta(s)", len(proposals))
        return []

    out: List[Dict] = []
    for p in proposals:
        iid = str(p.get("item_id") or "")
        acao = items_by_id.get(iid, "")
        is_completion = (p.get("action") == "complete") or (p.get("new_status") == "concluido")
        if acao and _is_judgment_item(acao) and is_completion:
            if not _references_deliverable(text, p.get("evidencia") or "", acao):
                logger.info(
                    "smart_updates: DROP conclusao de item de julgamento sem ref ao entregavel "
                    "(item=%r evid=%r)", acao[:50], (p.get("evidencia") or "")[:50]
                )
                continue
            if not _judgment_autoclose_enabled() and (p.get("confianca") or "").lower() == "alta":
                logger.info("smart_updates: rebaixa item de julgamento alta->media (propor, nao executar): %r", acao[:50])
                p = {**p, "confianca": "media"}
        out.append(p)
    return out


def _get_open_items(empresa_id: str) -> List[Dict]:
    """Retorna itens nao-concluidos da empresa pra usar como contexto do classifier."""
    import psycopg2, psycopg2.extras
    if not CONSELHOOS_DATABASE_URL:
        return []
    try:
        conn = psycopg2.connect(CONSELHOOS_DATABASE_URL)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.execute("""
            SELECT id, acao, status, prazo, responsavel_r
            FROM raci_itens
            WHERE empresa_id = %s AND status != 'concluido'
            ORDER BY prazo NULLS LAST
        """, (empresa_id,))
        rows = [dict(r) for r in cur.fetchall()]
        conn.close()
        return rows
    except Exception as e:
        logger.warning(f"_get_open_items({empresa_id}) falhou: {e}")
        return []


async def propose_updates_from_text(text: str, empresa_id: str) -> List[Dict]:
    """Le msg livre + contexto dos itens abertos da empresa. Retorna lista de
    propostas: [{item_id, acao_short, action, new_status, new_prazo, notes,
    evidencia, confianca}]. Vazio se nada matchou.

    confianca = 'alta' (auto-apply) | 'media' (notif Renato) | 'baixa' (ignora).
    """
    if not ANTHROPIC_API_KEY:
        return []
    text = (text or "").strip()
    if len(text) < MIN_TEXT_LEN_FOR_AI:
        return []
    items = _get_open_items(empresa_id)
    if not items:
        return []

    items_str = "\n".join([
        f"[{i+1}] id={it['id']} | {it['acao']}\n    Responsavel: {it.get('responsavel_r') or '?'} | Prazo: {it['prazo']} | Status: {it['status']}"
        for i, it in enumerate(items)
    ])

    prompt = f"""Voce e analista de governanca interpretando mensagens informais de grupo WhatsApp do Conselho.

CONTEXTO: itens RACI abertos (nao-concluidos) da empresa. Cada um tem um id UUID.

{items_str}

MENSAGEM RECEBIDA: \"\"\"{text}\"\"\"

TAREFA: identifique se a mensagem reporta status sobre algum dos itens acima. Pode mencionar 0, 1 ou varios itens.

Responda JSON array (vazio se nada matchou):

[
  {{
    "item_id": "<uuid exato>",
    "action": "update_status" | "update_prazo" | "add_note" | "complete",
    "new_status": "concluido" | "em_andamento" | "pendente" | "cancelado" | null,
    "new_prazo": "YYYY-MM-DD" | null,
    "notes": "<texto extraido da msg, max 200 chars>" | null,
    "evidencia": "<trecho exato da msg que justifica, max 100 chars>",
    "confianca": "alta" | "media" | "baixa"
  }}
]

Regras:
- IMPORTANTE: item_id DEVE ser exatamente um UUID listado acima. NUNCA invente ID, NUNCA use numero. Se nao achar item correspondente claro, NAO inclua a proposta.
- confianca=alta: msg cita item de forma inequivoca (acao, nome do responsavel, dados especificos). Ex: "Kommo CRM ativado hoje" + tem item "Ativar Kommo CRM".
- confianca=media: msg pode estar relacionada mas ambigua.
- confianca=baixa: matching forçado, contexto insuficiente. Prefira retornar vazio em vez de baixa.
- Mensagens irrelevantes (saudacao, off-topic, emoji isolado): retorne array vazio [].
- CORTESIA NAO E CONCLUSAO: mensagens sociais/gratidao pos-reuniao ("fico feliz em colaborar",
  "grato pelos feedbacks", "agradeco a reuniao", "conte comigo", "sucesso a todos") NAO fecham
  nem alteram item nenhum. Se a msg e so cortesia, retorne [].
- ITEM DE REVISAO/APROVACAO/AVALIACAO (verbos "revisar", "aprovar", "validar", "avaliar", "analisar"):
  so marque concluido se a msg comentar EXPLICITAMENTE O DOCUMENTO/ENTREGAVEL em si (ex: "revisei o
  questionario", "aprovei a minuta", "o documento esta ok"). Falar bem da reuniao ou agradecer NAO
  conta. Na duvida, use confianca "media" (vira proposta pro humano), nunca "alta".
- DETECCAO DE CONCLUSAO — marque new_status=concluido quando a msg indica:
  * Verbos no PASSADO ("apresentado", "fechado", "criado", "definido", "implementado", "realizado", "feito", "concluido")
  * Anuncio de RESULTADO final ("decisao tomada", "ficou em X", "esta pronto")
  * Acao especifica do RACI item executada (mesmo se sem palavra "concluido")
  * Ex: item "Conversar com Camila" + msg "Reuniao com Camila apresentado Clube Vallen e Politica Comercial" => concluido (a conversa aconteceu).
- Se reporta progresso sem finalizar ("estamos trabalhando", "em curso", "ainda falta X") → new_status=em_andamento.
- Se anuncia novo prazo → action=update_prazo + new_prazo.
- Se adiciona contexto sem mudar status → action=add_note.

Responda APENAS o JSON array."""

    try:
        async with httpx.AsyncClient(timeout=20.0) as cli:
            r = await cli.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                json={"model": llm.FAST, "max_tokens": 1200,
                      "messages": [{"role": "user", "content": prompt}]}
            )
        if r.status_code != 200:
            logger.warning(f"propose_updates_from_text: API {r.status_code}: {r.text[:200]}")
            return []
        _llm_resp = r.json()
        llm_usage.record_response("raci.update", llm.FAST, _llm_resp)  # F-E: custo por-funcao
        text_out = _llm_resp["content"][0]["text"]
        # Extrai JSON array
        start = text_out.find('[')
        end = text_out.rfind(']') + 1
        if start < 0:
            return []
        proposals = json.loads(text_out[start:end])
        # Validacao basica + filtro de alucinacao de IDs (UUID real do contexto)
        # Bug detectado 08/06/26: Claude as vezes inventa item_id="13" ou UUIDs ineditos.
        valid_ids = {str(it['id']) for it in items}
        import re
        UUID_RE = re.compile(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', re.I)
        valid = []
        for p in proposals:
            if not isinstance(p, dict): continue
            iid = str(p.get('item_id') or '')
            if not iid or not UUID_RE.match(iid):
                logger.info(f"smart_updates: drop proposta item_id invalido '{iid[:40]}'")
                continue
            if iid not in valid_ids:
                logger.info(f"smart_updates: drop proposta item_id alucinado '{iid}' (nao esta no contexto)")
                continue
            if not p.get('confianca'):
                continue
            valid.append(p)
        # Rede de seguranca deterministica: filtra cortesia + protege itens de julgamento.
        items_by_id = {str(it['id']): (it.get('acao') or '') for it in items}
        return _apply_guardrails(valid, items_by_id, text)
    except Exception as e:
        logger.warning(f"propose_updates_from_text falhou: {type(e).__name__}: {e}")
        return []


def apply_proposal(proposal: Dict, empresa_id: str) -> Optional[Dict]:
    """Aplica 1 proposta no DB + audit log. Retorna {acao, old_status, new_status}
    pra UI/confirmacao no grupo. None se item nao encontrado."""
    import psycopg2, psycopg2.extras
    if not CONSELHOOS_DATABASE_URL:
        return None
    try:
        conn = psycopg2.connect(CONSELHOOS_DATABASE_URL)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.execute("SELECT id, acao, status FROM raci_itens WHERE id = %s", (proposal['item_id'],))
        row = cur.fetchone()
        if not row:
            conn.close()
            return None
        target = dict(row)

        # Build update — skip se nao tem mudanca real (msg-pergunta ou contexto)
        sets = []
        params = []
        new_status = proposal.get('new_status')
        if new_status and new_status != target['status']:
            sets.append("status = %s"); params.append(new_status)
        if proposal.get('new_prazo'):
            sets.append("prazo = %s"); params.append(proposal['new_prazo'])
        if proposal.get('notes'):
            sets.append("notas = COALESCE(notas, '') || %s")
            from services.tz import now_utc, format_brt
            params.append(f"\n[{format_brt(now_utc(), '%d/%m')}] {proposal['notes'][:200]}")
        if not sets:
            conn.close()
            return None  # no-op — nao polui audit log nem reply no grupo
        sets.append("updated_at = NOW()")
        params.append(target['id'])
        cur.execute(f"UPDATE raci_itens SET {', '.join(sets)} WHERE id = %s", tuple(params))
        conn.commit()
        conn.close()

        # Audit log
        try:
            from services.agent_actions import log_action
            log_action(
                action_type='raci_status_updated',
                category='conselho',
                title=f"RACI (AI): '{(target['acao'] or '')[:60]}' → {proposal.get('new_status') or 'note'}",
                scope_ref={'raci_item_id': str(target['id']), 'empresa_id': str(empresa_id)},
                source='raci_smart_updates.apply_proposal',
                payload={
                    'old_status': target['status'],
                    'new_status': proposal.get('new_status'),
                    'new_prazo': proposal.get('new_prazo'),
                    'notes': proposal.get('notes'),
                    'evidencia': proposal.get('evidencia'),
                    'confianca': proposal.get('confianca'),
                },
                undo_hint=f"UPDATE raci_itens SET status='{target['status']}' WHERE id='{target['id']}'::uuid;",
            )
        except Exception as e:
            logger.warning(f"audit log apply_proposal falhou: {e}")

        return {
            'item_id': target['id'],
            'acao': target['acao'],
            'old_status': target['status'],
            'new_status': proposal.get('new_status') or target['status'],
            'evidencia': proposal.get('evidencia'),
        }
    except Exception as e:
        logger.warning(f"apply_proposal falhou: {e}")
        return None


async def _download_media_from_evolution(message_key: Dict, instance: str) -> Optional[Dict]:
    """Baixa media (base64 + mime) da Evolution API via key.id.
    Returns {'base64', 'mimetype'} ou None. Timeout 15s pra nao bloquear webhook."""
    evo_url = (os.getenv("EVOLUTION_API_URL") or "").rstrip('/')
    evo_key = (os.getenv("EVOLUTION_API_KEY") or "").strip()
    if not evo_url or not evo_key or not instance:
        return None
    try:
        async with httpx.AsyncClient(timeout=15.0) as cli:
            r = await cli.post(
                f"{evo_url}/chat/getBase64FromMediaMessage/{instance}",
                headers={"apikey": evo_key, "Content-Type": "application/json"},
                json={"message": {"key": message_key}, "convertToMp4": False},
            )
        if r.status_code not in (200, 201):
            logger.warning(f"download media: HTTP {r.status_code}")
            return None
        d = r.json()
        b64 = d.get("base64") or ""
        mime = (d.get("mimetype") or "").split(";")[0].strip()
        if not b64:
            return None
        if len(b64) * 0.75 > MAX_MEDIA_BYTES:
            logger.warning(f"media too large: ~{int(len(b64)*0.75/1024/1024)}MB")
            return None
        return {"base64": b64, "mimetype": mime or "application/octet-stream"}
    except Exception as e:
        logger.warning(f"download media error: {e}")
        return None


async def _claude_media_to_text(b64: str, mime: str, instruction: str,
                                  model: str = llm.FAST) -> Optional[str]:
    """Envia media pro Claude com instrucao de extracao. Audio/PDF = type=document; image = type=image."""
    if not ANTHROPIC_API_KEY:
        return None
    is_image = mime.startswith("image/")
    block = (
        {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}}
        if is_image
        else {"type": "document", "source": {"type": "base64", "media_type": mime, "data": b64}}
    )
    try:
        async with httpx.AsyncClient(timeout=60.0) as cli:
            r = await cli.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                         "content-type": "application/json"},
                json={"model": model, "max_tokens": 2000,
                      "messages": [{"role": "user", "content": [block, {"type": "text", "text": instruction}]}]},
            )
        if r.status_code != 200:
            logger.warning(f"claude media: HTTP {r.status_code}: {r.text[:200]}")
            return None
        _llm_resp = r.json()
        llm_usage.record_response("raci.smart_update", model, _llm_resp)  # F-E: custo por-funcao
        return _llm_resp["content"][0]["text"].strip()
    except Exception as e:
        logger.warning(f"claude media error: {e}")
        return None


def _extract_docx_text(b64: str) -> Optional[str]:
    """Extrai texto de .docx localmente via python-docx (paragrafos + tabelas)."""
    try:
        import base64, io
        from docx import Document
        doc = Document(io.BytesIO(base64.b64decode(b64)))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += "\n" + cell.text
        return text or None
    except Exception as e:
        logger.warning(f"docx extract error: {e}")
        return None


async def extract_text_from_media(message: Dict, message_key: Dict, instance: str,
                                    caption: str = "") -> Optional[str]:
    """Detecta tipo de media (audio/image/document) e extrai texto via Claude
    multimodal ou python-docx. Junta caption se houver. Returns text ou None."""
    has_audio = "audioMessage" in message
    has_image = "imageMessage" in message
    has_doc = "documentMessage" in message
    if not (has_audio or has_image or has_doc):
        return None

    media = await _download_media_from_evolution(message_key, instance)
    if not media:
        return None
    b64 = media["base64"]
    mime = media["mimetype"]

    extracted: Optional[str] = None

    if has_audio:
        if not mime or mime == "application/octet-stream":
            mime = "audio/ogg"  # WA default
        extracted = await _claude_media_to_text(
            b64, mime,
            "Transcreva este audio em portugues. Responda APENAS a transcricao, sem comentarios."
        )

    elif has_image:
        if not mime.startswith("image/"):
            mime = "image/jpeg"
        extracted = await _claude_media_to_text(
            b64, mime,
            "Descreva o conteudo desta imagem em portugues — texto, dados, contexto. "
            "Se for screenshot de planilha/dashboard/conversa, transcreva o conteudo literal. "
            "Maximo 500 palavras."
        )

    elif has_doc:
        doc_msg = message.get("documentMessage", {}) or {}
        filename = (doc_msg.get("fileName") or "").lower()
        if filename.endswith(".docx") or "wordprocessingml" in mime:
            extracted = _extract_docx_text(b64)
        elif filename.endswith(".pdf") or "pdf" in mime:
            extracted = await _claude_media_to_text(
                b64, "application/pdf",
                "Extraia o conteudo textual deste PDF em portugues. Foque em dados, decisoes, "
                "prazos, nomes. Pule cabecalhos/rodapes repetitivos. Maximo 2000 palavras."
            )
        elif filename.endswith((".txt", ".md")) or mime.startswith("text/"):
            try:
                import base64
                extracted = base64.b64decode(b64).decode("utf-8", errors="replace")[:5000]
            except Exception:
                pass
        else:
            logger.info(f"documentMessage tipo nao suportado: {filename} mime={mime}")

    if not extracted:
        return None

    extracted = extracted.strip()
    if caption and caption.strip():
        extracted = f"[Legenda: {caption.strip()}]\n\n{extracted}"
    return extracted


async def process_group_message(text: str, empresa_id: str, empresa_nome: str = "") -> Dict:
    """Pipeline completo: tenta regex, fallback AI, aplica alta confianca,
    notifica Renato pra media confianca. Retorna resumo do que rolou.

    Returns dict com:
      - applied: list de itens atualizados (pra confirmar no grupo)
      - pending_review: list de propostas que precisam aprovacao Renato
      - skipped_low: count de propostas baixa-conf descartadas
    """
    # 1. Fast-path: regex (free)
    try:
        from services.raci_weekly_report import parse_raci_update
        regex_result = parse_raci_update(text, empresa_id)
        if regex_result:
            return {"applied": [regex_result], "pending_review": [], "skipped_low": 0, "source": "regex"}
    except Exception as e:
        logger.warning(f"regex fast-path falhou: {e}")

    # 2. AI fallback
    proposals = await propose_updates_from_text(text, empresa_id)
    if not proposals:
        return {"applied": [], "pending_review": [], "skipped_low": 0, "source": "ai_no_match"}

    applied = []
    pending = []
    skipped = 0
    for p in proposals:
        conf = (p.get('confianca') or '').lower()
        if conf == 'alta':
            r = apply_proposal(p, empresa_id)
            if r: applied.append(r)
        elif conf == 'media':
            pending.append(p)
        else:
            skipped += 1

    # 3. Notif Renato pra media
    if pending:
        try:
            from services.intel_bot import send_intel_notification
            lines = [f"⚠️ Update RACI {empresa_nome or 'empresa'} precisa aprovacao:"]
            for p in pending:
                lines.append(f"- _{(p.get('evidencia') or '')[:80]}_")
                lines.append(f"  → propor {p.get('new_status') or 'note'} (conf {p.get('confianca')})")
            lines.append(f"\nResponda no app /editorial ou edite manual no ConselhoOS.")
            await send_intel_notification("\n".join(lines))
        except Exception as e:
            logger.warning(f"notif Renato pending review falhou: {e}")

    return {"applied": applied, "pending_review": pending, "skipped_low": skipped, "source": "ai"}


# ============== BATCH PROCESSOR (08/jun/2026) ==============

async def process_week_for_empresa(
    empresa_id: str,
    group_jid: str,
    empresa_nome: str = "",
    days: int = 7,
    auto_apply: bool = False,
) -> Dict:
    """Roda smart_updates em batch sobre todas msgs de texto do grupo nos ultimos N dias.

    Args:
        empresa_id: UUID da empresa no ConselhoOS
        group_jid: JID do grupo WA (ex: '120363408325592607@g.us')
        empresa_nome: nome amigavel pra logs
        days: janela retroativa (default 7)
        auto_apply: se True, aplica propostas alta-confianca direto

    Returns:
        {
          'msgs_processed': int,
          'proposals_all': list[dict] (todas propostas geradas, ja deduped + match items reais),
          'high_confidence': list[dict],
          'medium_confidence': list[dict],
          'applied': list[dict] (se auto_apply=True),
        }
    """
    from database import get_db
    out: Dict = {
        "empresa_id": empresa_id,
        "empresa_nome": empresa_nome,
        "group_jid": group_jid,
        "msgs_processed": 0,
        "proposals_all": [],
        "high_confidence": [],
        "medium_confidence": [],
        "applied": [],
    }

    # Pega msgs de texto recentes, mais antigas primeiro pra Claude entender contexto cronologico
    with get_db() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            SELECT sender_name, content, timestamp, from_me
            FROM group_messages
            WHERE group_jid = %s
              AND timestamp > NOW() - (%s || ' days')::interval
              AND content IS NOT NULL AND length(content) >= %s
            ORDER BY timestamp ASC
            """,
            (group_jid, str(days), MIN_TEXT_LEN_FOR_AI),
        )
        msgs = cur.fetchall()

    out["msgs_processed"] = len(msgs)
    if not msgs:
        return out

    # Dedup propostas por item_id (mantem a de maior confianca)
    best_per_item: Dict[str, Dict] = {}

    for m in msgs:
        sender = "EU" if m.get('from_me') else (m.get('sender_name') or '?')
        ts = m.get('timestamp')
        text = m.get('content') or ''
        # Prefix com sender pra Claude ter contexto de quem fala
        prefixed = f"[{sender} em {ts.strftime('%d/%m %H:%M') if ts else '?'}]\n{text}"
        try:
            props = await propose_updates_from_text(prefixed, empresa_id)
        except Exception as e:
            logger.warning(f"process_week msg falhou: {e}")
            continue
        for p in props:
            iid = p['item_id']
            existing = best_per_item.get(iid)
            conf_rank = {'alta': 3, 'media': 2, 'baixa': 1}
            new_rank = conf_rank.get(p.get('confianca', '').lower(), 0)
            old_rank = conf_rank.get((existing or {}).get('confianca', '').lower(), 0)
            if new_rank > old_rank:
                best_per_item[iid] = p

    all_proposals = list(best_per_item.values())
    out["proposals_all"] = all_proposals
    out["high_confidence"] = [p for p in all_proposals if (p.get('confianca') or '').lower() == 'alta']
    out["medium_confidence"] = [p for p in all_proposals if (p.get('confianca') or '').lower() == 'media']

    if auto_apply:
        for p in out["high_confidence"]:
            r = apply_proposal(p, empresa_id)
            if r:
                out["applied"].append(r)

    return out
