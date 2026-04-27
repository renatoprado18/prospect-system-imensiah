"""
Gerador de Ata de Conselho em formato DOCX profissional.

TEMPLATE PADRAO — usado quando a empresa nao tem template especifico.

Formato:
- Header centralizado com nome da empresa
- Metadata compacta (4 colunas: duracao, participantes, itens, proxima)
- Participantes como cards (3 colunas, nome bold + papel cinza)
- Secoes numeradas com subsecoes
- Tabelas financeiras com header escuro e zebra
- Matriz RACI com iniciais (TM, GC, etc.) e areas abreviadas (Mkt, Fin, etc.)
- Pendencias com emojis coloridos (vermelho/amarelo/verde)
- Proximos passos como tabela 2 colunas
- Rodape com proxima reuniao + confidencial

Uso standalone:
    python scripts/ata_to_docx.py

Uso como modulo:
    from scripts.ata_to_docx import generate_ata_docx
    generate_ata_docx(ata_md, "EMPRESA", "08 de Abril de 2026", "/tmp/output.docx")

Para templates especificos por empresa, crie um modulo em scripts/templates/
que exporte a mesma funcao generate_ata_docx(ata_md, empresa_nome, data_reuniao, output_path).
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Color palette
# ---------------------------------------------------------------------------
COLOR_DARK = RGBColor(0x2C, 0x3E, 0x50)
COLOR_SUBTITLE = RGBColor(0x7F, 0x8C, 0x8D)
COLOR_DATE = RGBColor(0x95, 0xA5, 0xA6)
COLOR_GRAY = RGBColor(0x95, 0xA5, 0xA6)
COLOR_LIGHT_GRAY = RGBColor(0xBD, 0xBD, 0xBD)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_RED = RGBColor(0xC0, 0x39, 0x2B)
COLOR_YELLOW = RGBColor(0x9A, 0x7D, 0x0A)
COLOR_GREEN = RGBColor(0x27, 0xAE, 0x60)
COLOR_TABLE_HEADER_BG = "2c3e50"
COLOR_TABLE_ALT_BG = "f8f9fa"
COLOR_RACI_R = "d6eaf8"
COLOR_RACI_A = "d5f5e3"
COLOR_RACI_C = "fef9e7"
COLOR_RACI_I = "f2f3f4"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right, each a dict with sz, color, val."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    for edge, props in kwargs.items():
        edge_el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{props.get("val", "single")}" '
            f'w:sz="{props.get("sz", "4")}" w:space="0" '
            f'w:color="{props.get("color", "d0d0d0")}"/>'
        )
        tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def _remove_cell_borders(cell):
    """Remove all borders from a cell."""
    _set_cell_border(cell,
        top={"val": "none", "sz": "0", "color": "FFFFFF"},
        bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
        left={"val": "none", "sz": "0", "color": "FFFFFF"},
        right={"val": "none", "sz": "0", "color": "FFFFFF"},
    )


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in EMU."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def _add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name="Arial"):
    """Add a run with formatting to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def _add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="bdc3c7"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_rich_text(paragraph, text, default_size=10, default_color=None, default_font="Arial"):
    """Parse markdown bold (**text**) within a line and add runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            _add_run(paragraph, part[2:-2], bold=True, size=default_size,
                     color=default_color, font_name=default_font)
        else:
            _add_run(paragraph, part, size=default_size,
                     color=default_color, font_name=default_font)


# ---------------------------------------------------------------------------
# Markdown preprocessor — normalizes Claude output to expected format
# ---------------------------------------------------------------------------

def _preprocess_markdown(md: str) -> str:
    """Convert Claude's markdown ata to the format parse_ata_markdown expects.

    Handles:
    - # Title → plain text (empresa name)
    - ## SECTION → plain text with section markers
    - **bold** → keeps as-is (handled by _add_formatted_text)
    - | tables | → keeps as-is (handled by section body renderer)
    - ### subsections → numbered subsections
    """
    lines = md.split('\n')
    out = []
    section_num = 0
    empresa = ''
    found_header = False

    for line in lines:
        stripped = line.strip()

        # Skip empty markdown artifacts
        if stripped == '---':
            continue

        # # Title — extract empresa name
        if stripped.startswith('# ') and not found_header:
            found_header = True
            # Extract empresa from "# ALBA CONSULTORIA — Ata de Reunião de Conselho"
            title = stripped[2:]
            if '—' in title:
                empresa = title.split('—')[0].strip()
            elif '-' in title and 'Ata' in title:
                empresa = title.split('-')[0].strip().split('–')[0].strip()
            else:
                empresa = title
            out.append(empresa)
            out.append('')
            out.append('Ata de Reunião de Conselho')
            continue

        # **Data:** ... | **Duração:** ... — metadata line
        if stripped.startswith('**Data:**'):
            # Parse: **Data:** 16 de abril de 2026 | **Duração:** ~1h | **Participantes:** 6
            parts = stripped.split('|')
            date_part = parts[0].replace('**Data:**', '').strip() if parts else ''
            out.append(f"{date_part} · Reunião ordinária mensal")
            out.append('')
            # Extract metadata
            meta = {}
            for part in parts:
                part = part.strip().replace('**', '')
                if 'Duração:' in part:
                    meta['duracao'] = part.split('Duração:')[1].strip()
                elif 'Participantes:' in part:
                    meta['participantes'] = part.split('Participantes:')[1].strip()
            out.append('Duração')
            out.append(meta.get('duracao', '—'))
            out.append('')
            out.append('Participantes')
            out.append(meta.get('participantes', '—'))
            out.append('')
            out.append('Itens de ação')
            out.append('—')
            out.append('')
            out.append('Próxima reunião')
            out.append('—')
            out.append('')
            continue

        # ## PARTICIPANTES — section header
        if stripped.startswith('## PARTICIPANTES') or stripped == '## Participantes':
            out.append('PARTICIPANTES')
            out.append('')
            continue

        # ## DECISÕES APROVADAS
        if 'DECISÕES' in stripped.upper() and stripped.startswith('##'):
            title = stripped.lstrip('#').strip()
            section_num += 1
            out.append(f'{section_num}. {title}')
            out.append('')
            continue

        # ## PENDÊNCIAS
        if 'PENDÊNCIAS' in stripped.upper() and stripped.startswith('##'):
            title = stripped.lstrip('#').strip()
            section_num += 1
            out.append(f'{section_num}. {title}')
            out.append('')
            continue

        # ## PRÓXIMA REUNIÃO
        if 'PRÓXIMA REUNIÃO' in stripped.upper() and stripped.startswith('##'):
            title = stripped.lstrip('#').strip()
            section_num += 1
            out.append(f'{section_num}. {title}')
            out.append('')
            continue

        # ## N. SECTION TITLE — numbered section
        m = re.match(r'^##\s+(\d+)\.\s+(.+)$', stripped)
        if m:
            section_num = int(m.group(1))
            out.append(f'{section_num}. {m.group(2)}')
            out.append('')
            continue

        # ## SECTION TITLE — unnumbered section (make it numbered)
        if stripped.startswith('## '):
            title = stripped[3:].strip()
            section_num += 1
            out.append(f'{section_num}. {title}')
            out.append('')
            continue

        # ### Subsection
        m = re.match(r'^###\s+(\d+\.\d+)\s+(.+)$', stripped)
        if m:
            out.append(f'{m.group(1)} {m.group(2)}')
            continue

        if stripped.startswith('### '):
            title = stripped[4:].strip()
            out.append(title)
            continue

        # Regular line — pass through (keeps **bold**, tables, bullets)
        out.append(line)

    return '\n'.join(out)


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def parse_ata_markdown(md: str) -> dict:
    """Parse the ata markdown into structured sections."""
    lines = md.split('\n')
    result = {
        'empresa': '',
        'subtitulo': '',
        'data_linha': '',
        'metadata': {},  # duracao, participantes, itens_acao, proxima
        'participantes_text': '',
        'sections': [],
        'footer_date': '',
    }

    # Extract header info (first few lines)
    i = 0
    while i < len(lines) and i < 10:
        line = lines[i].strip()
        if line and not result['empresa']:
            result['empresa'] = line
            i += 1
            continue
        if 'Ata de Reunião' in line:
            result['subtitulo'] = line
            i += 1
            continue
        if 'Reunião ordinária' in line or ('·' in line and '2026' in line):
            result['data_linha'] = line
            i += 1
            continue
        i += 1

    # Extract metadata block (Duração, Participantes, Itens de ação, Próxima reunião)
    # The format is: 4 label lines, then 4 value groups separated by blank lines
    metadata_labels = {'Duração', 'Participantes', 'Itens de ação', 'Próxima reunião'}

    # Find where labels end and values begin
    meta_start = None
    meta_end = None
    for idx, line in enumerate(lines):
        if line.strip() == 'Duração':
            meta_start = idx
        if meta_start is not None and line.strip() == 'PARTICIPANTES':
            meta_end = idx
            break

    if meta_start is not None and meta_end is not None:
        # Collect non-label, non-empty lines between labels section and PARTICIPANTES
        # First skip the 4 label lines, then collect value groups
        value_lines = []
        past_labels = False
        for idx in range(meta_start, meta_end):
            stripped = lines[idx].strip()
            if stripped in metadata_labels:
                continue
            value_lines.append(stripped)

        # Group by blank line separators
        groups = []
        current = []
        for vl in value_lines:
            if vl == '':
                if current:
                    groups.append(' '.join(current))
                    current = []
            else:
                current.append(vl)
        if current:
            groups.append(' '.join(current))

        if len(groups) >= 1:
            result['metadata']['duracao'] = groups[0]
        if len(groups) >= 2:
            result['metadata']['participantes'] = groups[1]
        if len(groups) >= 3:
            result['metadata']['itens_acao'] = groups[2]
        if len(groups) >= 4:
            result['metadata']['proxima'] = groups[3]

    # Extract participants
    part_match = re.search(r'PARTICIPANTES\s*\n\s*\n(.*?)(?=\n\s*\n\s*\d+\.)', md, re.DOTALL)
    if part_match:
        result['participantes_text'] = part_match.group(1).strip()

    # Extract sections (numbered: 1., 2., etc.)
    # Split by numbered section headers
    section_pattern = re.compile(r'^(\d+)\.\s+(.+)$', re.MULTILINE)
    matches = list(section_pattern.finditer(md))

    for idx, match in enumerate(matches):
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(md)
        section_body = md[start:end].strip()
        result['sections'].append({
            'number': match.group(1),
            'title': match.group(2).strip(),
            'body': section_body,
        })

    # Footer date
    proxima_match = re.search(r'próxima reunião ordinária.*?(\d+\s+de\s+\w+\s+de\s+\d{4})', md, re.IGNORECASE)
    if proxima_match:
        result['footer_date'] = proxima_match.group(1)
    elif result['metadata'].get('proxima'):
        result['footer_date'] = result['metadata']['proxima']

    return result


def _parse_pipe_table(text: str) -> list[list[str]]:
    """Parse a markdown pipe table into rows of cells."""
    rows = []
    for line in text.strip().split('\n'):
        line = line.strip()
        if not line.startswith('|'):
            continue
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue  # separator row
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def _extract_tables_and_text(body: str) -> list:
    """Split section body into text blocks and table blocks."""
    blocks = []
    lines = body.split('\n')
    current_text = []
    current_table = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        is_table_line = stripped.startswith('|') and '|' in stripped[1:]
        is_separator = bool(re.match(r'^\|[\s\-:|]+\|$', stripped))

        if is_table_line:
            if not in_table:
                # Flush text
                if current_text:
                    blocks.append(('text', '\n'.join(current_text)))
                    current_text = []
                in_table = True
            current_table.append(stripped)
        else:
            if in_table:
                # Flush table
                if current_table:
                    blocks.append(('table', '\n'.join(current_table)))
                    current_table = []
                in_table = False
            current_text.append(line)

    # Flush remaining
    if current_table:
        blocks.append(('table', '\n'.join(current_table)))
    if current_text:
        text = '\n'.join(current_text).strip()
        if text:
            blocks.append(('text', text))

    return blocks


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def generate_ata_docx(ata_md: str, empresa_nome: str, data_reuniao: str, output_path: str) -> str:
    """
    Generate a professionally formatted DOCX from ata markdown.

    Args:
        ata_md: The ata content in markdown format
        empresa_nome: Company name (e.g. "VALLEN CLINIC")
        data_reuniao: Meeting date string (e.g. "08 de Abril de 2026")
        output_path: Path for the output DOCX file

    Returns:
        output_path
    """
    # Preprocess markdown from Claude to expected format
    ata_md = _preprocess_markdown(ata_md)
    parsed = parse_ata_markdown(ata_md)
    doc = Document()

    # -----------------------------------------------------------------------
    # Page setup: A4, portrait, margins
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # -----------------------------------------------------------------------
    # Default font
    # -----------------------------------------------------------------------
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 1.15

    # -----------------------------------------------------------------------
    # Header block (centered)
    # -----------------------------------------------------------------------
    # Company name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=2)
    _add_run(p, empresa_nome, bold=True, size=18, color=COLOR_DARK)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=2)
    _add_run(p, "Ata de Reunião de Conselho", size=12, color=COLOR_SUBTITLE)

    # Date line
    date_line = parsed.get('data_linha', '') or f"{data_reuniao} · Reunião ordinária mensal"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=0, after=6)
    _add_run(p, date_line, size=11, color=COLOR_DATE)

    # Horizontal rule
    _add_horizontal_rule(doc)

    # -----------------------------------------------------------------------
    # Metadata table (1 row, 4 columns, no borders)
    # -----------------------------------------------------------------------
    meta = parsed.get('metadata', {})
    meta_labels = ['Duração', 'Participantes', 'Itens de ação', 'Próxima reunião']
    meta_values = [
        meta.get('duracao', '—'),
        meta.get('participantes', '—'),
        meta.get('itens_acao', '—'),
        meta.get('proxima', '—'),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="FFFFFF"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for col_idx, (label, value) in enumerate(zip(meta_labels, meta_values)):
        cell = table.cell(0, col_idx)
        _remove_cell_borders(cell)
        _set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
        # Clear default paragraph
        cell.text = ''
        # Label (small caps style - uppercase + small font)
        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p_label, before=0, after=1)
        _add_run(p_label, label.upper(), size=7, color=COLOR_GRAY)
        # Value
        p_val = cell.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(p_val, before=1, after=0)
        _add_run(p_val, value, bold=True, size=10, color=COLOR_DARK)

    # Spacing after metadata
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Participants as formatted table
    # -----------------------------------------------------------------------
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=6, after=4)
    _add_run(p, "PARTICIPANTES", bold=True, size=9, color=COLOR_SUBTITLE)

    if parsed['participantes_text']:
        # Parse participants: "Name (Role) | Name (Role) | ..."
        import re as _re
        raw = parsed['participantes_text'].replace('\n', ' ').replace('  ', ' ')
        # Split by | but also handle cases where | is missing (line breaks became spaces)
        # First split by |
        parts = [p.strip() for p in raw.split('|') if p.strip()]
        # Then split any part that has multiple "Name (Role)" patterns
        expanded = []
        for part in parts:
            # Match "Name (Role) Name2 (Role2)" → split into separate entries
            sub_parts = _re.findall(r'[A-ZÀ-Ý][^()]*\([^)]+\)', part)
            if len(sub_parts) > 1:
                expanded.extend([sp.strip() for sp in sub_parts])
            else:
                expanded.append(part)
        parts = expanded

        # Full name mappings (markdown may use short names)
        full_names = {
            'guilherme': 'Guilherme Cintra',
            'renato': 'Renato de Faria e Almeida Prado',
        }

        participants = []
        for part in parts:
            m = _re.match(r'^(.+?)\s*\((.+?)\)\s*$', part)
            if m:
                name = m.group(1).strip()
                role = m.group(2).strip()
                # Replace short names with full names
                name_lower = name.lower().strip()
                if name_lower in full_names:
                    name = full_names[name_lower]
                participants.append((name, role))
            else:
                name = part.strip()
                name_lower = name.lower()
                if name_lower in full_names:
                    name = full_names[name_lower]
                participants.append((name, ''))

        # Create table: 3 columns, as many rows as needed
        cols = 3
        rows_needed = (len(participants) + cols - 1) // cols
        table = doc.add_table(rows=rows_needed, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # No borders, light background
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="none" w:sz="0" w:space="0"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tblPr.append(tblW)

        for idx, (name, role) in enumerate(participants):
            row_idx = idx // cols
            col_idx = idx % cols
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            _set_cell_shading(cell, "f0f4f8")
            _set_cell_margins(cell, top=40, bottom=40, left=60, right=60)
            cp = cell.paragraphs[0]
            _set_paragraph_spacing(cp, before=0, after=0, line_spacing=1.1)
            _add_run(cp, name, bold=True, size=9, color=COLOR_DARK)
            if role:
                cp.add_run('\n')
                _add_run(cp, role, size=7, color=COLOR_GRAY)

        # Fill empty cells
        for idx in range(len(participants), rows_needed * cols):
            row_idx = idx // cols
            col_idx = idx % cols
            cell = table.cell(row_idx, col_idx)
            cell.text = ''

    _add_horizontal_rule(doc)

    # -----------------------------------------------------------------------
    # Sections
    # -----------------------------------------------------------------------
    for section_data in parsed['sections']:
        _render_section(doc, section_data)

    # -----------------------------------------------------------------------
    # Footer text
    # -----------------------------------------------------------------------
    footer_date = parsed.get('footer_date', '') or data_reuniao
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=20, after=4)
    _add_horizontal_rule(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, before=4, after=0)
    _add_run(p, f"Próxima reunião de conselho: {footer_date} | Confidencial",
             size=9, italic=True, color=COLOR_GRAY)

    # -----------------------------------------------------------------------
    # Add running footer to section
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(fp, f"Próxima reunião de conselho: {footer_date} | Confidencial",
             size=8, italic=True, color=COLOR_GRAY)

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    doc.save(output_path)
    return output_path


def _render_section(doc, section_data: dict):
    """Render a numbered section with its body content."""
    number = section_data['number']
    title = section_data['title']
    body = section_data['body']

    # Section title
    p = doc.add_paragraph()
    _set_paragraph_spacing(p, before=18, after=6)
    _add_run(p, f"{number}. {title}", bold=True, size=13, color=COLOR_DARK)

    # Parse body into blocks
    blocks = _extract_tables_and_text(body)

    for block_type, content in blocks:
        if block_type == 'table':
            _render_table(doc, content, section_data.get('title', ''))
        else:
            _render_text_block(doc, content)


def _render_text_block(doc, text: str):
    """Render a text block with subsections, bullets, pendencias, etc."""
    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Skip footer lines repeated in source
        if line.startswith('Próxima reunião de conselho:') or re.match(r'^.*\| Ata do Conselho \|', line):
            i += 1
            continue

        # Subsection (e.g., "2.1 Title" or "**2.1 Title**")
        subsection_match = re.match(r'^\*?\*?(\d+\.\d+)\s+(.+?)\*?\*?\s*$', line)
        if subsection_match:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=12, after=4)
            sub_title = subsection_match.group(2).strip().rstrip('*')
            _add_run(p, f"{subsection_match.group(1)} {sub_title}",
                     bold=True, size=11, color=COLOR_DARK)
            i += 1
            continue

        # Pendencias critical/important/governance headers
        if line.startswith('🔴') or line.startswith('🟡') or line.startswith('🟢'):
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=10, after=4)
            # Determine color
            if '🔴' in line:
                color = COLOR_RED
            elif '🟡' in line:
                color = COLOR_YELLOW
            else:
                color = COLOR_GREEN
            # Clean the line
            clean_line = line.replace('🔴', '').replace('🟡', '').replace('🟢', '').strip()
            # Add emoji + text
            emoji = line[0:2] if line[0] in '🔴🟡🟢' else line[0]
            # Find the actual emoji
            for em in ['🔴', '🟡', '🟢']:
                if em in line:
                    emoji = em
                    break
            run = _add_run(p, f"{emoji} ", bold=True, size=11, color=color)
            _add_rich_text(p, clean_line, default_size=11, default_color=color)
            i += 1
            continue

        # Pendencia item (P1, P2, etc.)
        pend_match = re.match(r'^\*?\*?(P\d+)\s*·\s*(.+?)\*?\*?\s*$', line)
        if not pend_match:
            # Also match lines starting with **P1
            pend_match = re.match(r'^\*\*(P\d+)\s*·\s*(.+?)\*\*\s*$', line)
        if pend_match:
            # Collect continuation lines
            full_text = line
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('P') \
                    and not lines[i].strip().startswith('🔴') \
                    and not lines[i].strip().startswith('🟡') \
                    and not lines[i].strip().startswith('🟢') \
                    and not re.match(r'^\d+\.', lines[i].strip()) \
                    and not lines[i].strip().startswith('|') \
                    and not lines[i].strip().startswith('- ') \
                    and not lines[i].strip().startswith('* ') \
                    and not re.match(r'^\*\*P\d+', lines[i].strip()):
                full_text += ' ' + lines[i].strip()
                i += 1
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=4)
            p.paragraph_format.left_indent = Cm(0.5)
            _add_rich_text(p, full_text, default_size=10, default_color=COLOR_DARK)
            continue

        # Bullet points (- or * or •)
        bullet_match = re.match(r'^[-*•]\s+(.+)$', line)
        if bullet_match:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=2)
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.4)
            _add_run(p, "\u2022  ", size=10, color=COLOR_DARK)
            _add_rich_text(p, bullet_match.group(1), default_size=10, default_color=COLOR_DARK)
            i += 1
            continue

        # Regular paragraph
        # Collect continuation lines
        full_text = line
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                break
            if next_line.startswith('-') or next_line.startswith('*') or next_line.startswith('•'):
                break
            if re.match(r'^\d+\.\d+\s', next_line):
                break
            if re.match(r'^\d+\.\s', next_line):
                break
            if next_line.startswith('|'):
                break
            if next_line.startswith('🔴') or next_line.startswith('🟡') or next_line.startswith('🟢'):
                break
            if re.match(r'^\*\*P\d+', next_line) or re.match(r'^P\d+\s*·', next_line):
                break
            full_text += ' ' + next_line
            i += 1

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_paragraph_spacing(p, before=3, after=3)
        _add_rich_text(p, full_text, default_size=10, default_color=COLOR_DARK)


def _render_table(doc, table_text: str, section_title: str = ''):
    """Render a markdown pipe table as a formatted DOCX table."""
    rows = _parse_pipe_table(table_text)
    if not rows:
        return

    is_raci = 'RACI' in section_title.upper() or any(
        'Thalita' in cell or 'Gui' in cell or 'Amadeo' in cell
        for row in rows[:2] for cell in row
    )

    num_cols = max(len(row) for row in rows)
    # Ensure all rows have same number of columns
    for row in rows:
        while len(row) < num_cols:
            row.append('')

    # RACI: abbreviate headers and area column, track for legend
    raci_name_legend = {}   # abbrev → full name
    raci_area_legend = {}   # abbrev → full area name
    if is_raci:
        # Use initials (first letter of first + last name) when possible, else 3 letters
        name_abbrev = {
            'thalita': ('TM', 'Thalita Mendes'),
            'gui': ('GC', 'Guilherme Cintra'),
            'amadeo': ('AC', 'Amadeo Comin'),
            'renata': ('RC', 'Renata Comin'),
            'verid.': ('Ve', 'Veridiana'),
            'verid': ('Ve', 'Veridiana'),
            'lara': ('LG', 'Lara Gasparini'),
            'renato': ('RP', 'Renato Prado'),
        }
        area_abbrev = {
            'marketing': ('Mkt', 'Marketing'),
            'financeiro': ('Fin', 'Financeiro'),
            'rh': ('RH', 'Recursos Humanos'),
            'eq. médica': ('Eq.Med', 'Equipe Médica'),
            'eq.médica': ('Eq.Med', 'Equipe Médica'),
            'eq. medica': ('Eq.Med', 'Equipe Médica'),
            'operações': ('Oper', 'Operações'),
            'operacoes': ('Oper', 'Operações'),
            'ti / crm': ('TI', 'TI / CRM'),
            'ti/crm': ('TI', 'TI / CRM'),
            'ti': ('TI', 'TI / CRM'),
            'influencers': ('Infl', 'Influencers'),
            'estética': ('Est', 'Estética'),
            'estetica': ('Est', 'Estética'),
        }
        # Abbreviate header names (person columns)
        if rows:
            for ci in range(3, len(rows[0])):
                key = rows[0][ci].strip().lower()
                entry = name_abbrev.get(key)
                if entry:
                    rows[0][ci] = entry[0]
                    raci_name_legend[entry[0]] = entry[1]
        # Abbreviate area column in data rows (strip ** bold markers first)
        for ri in range(1, len(rows)):
            if rows[ri]:
                key = rows[ri][0].strip().replace('**', '').strip().lower()
                entry = area_abbrev.get(key)
                if entry:
                    rows[ri][0] = entry[0]
                    raci_area_legend[entry[0]] = entry[1]

    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # RACI: set aggressive column widths via XML
    if is_raci and num_cols >= 4:
        from docx.oxml.ns import qn
        tbl_xml = table._tbl
        # Remove auto column width and set fixed
        tblPr_temp = tbl_xml.tblPr if tbl_xml.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tblPr_temp.append(layout)

        # Set gridCol widths in twips (1 cm = 567 twips)
        tblGrid = tbl_xml.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gridCol in tblGrid.findall(qn('w:gridCol')):
                tblGrid.remove(gridCol)
        else:
            tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}/>')
            tbl_xml.insert(1, tblGrid)

        person_count = num_cols - 3
        # Total page width ~16cm = 9072 twips
        area_w = 800      # ~1.4cm
        acao_w = 9072 - 800 - 740 - (person_count * 400)  # remainder for Ação
        prazo_w = 740      # ~1.3cm
        person_w = 400     # ~0.7cm

        widths = [area_w, acao_w, prazo_w] + [person_w] * person_count
        for w in widths:
            gridCol = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{w}"/>')
            tblGrid.append(gridCol)

    # Set table style with borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="d0d0d0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="d0d0d0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="d0d0d0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="d0d0d0"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="d0d0d0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="d0d0d0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Set table width to full page width
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            _set_paragraph_spacing(p, before=1, after=1, line_spacing=1.0)
            _set_cell_margins(cell, top=30, bottom=30, left=50, right=50)

            clean_text = cell_text.strip().replace('**', '')

            if row_idx == 0:
                # Header row
                _set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _add_run(p, clean_text, bold=True, size=8, color=COLOR_WHITE)
            else:
                # Data rows
                if row_idx % 2 == 0:
                    _set_cell_shading(cell, COLOR_TABLE_ALT_BG)

                # RACI coloring (person columns start at 3)
                if is_raci and col_idx >= 3 and clean_text in ('R', 'A', 'C', 'I'):
                    raci_colors = {
                        'R': COLOR_RACI_R,
                        'A': COLOR_RACI_A,
                        'C': COLOR_RACI_C,
                        'I': COLOR_RACI_I,
                    }
                    _set_cell_shading(cell, raci_colors.get(clean_text, COLOR_TABLE_ALT_BG))
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(p, clean_text, bold=True, size=8, color=COLOR_DARK)
                elif is_raci and col_idx >= 3:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(p, clean_text, size=8, color=COLOR_GRAY)
                elif col_idx == 0:
                    # First column - left aligned, possibly bold
                    _add_rich_text(p, cell_text.strip(), default_size=9, default_color=COLOR_DARK)
                else:
                    _add_rich_text(p, cell_text.strip(), default_size=9, default_color=COLOR_DARK)

    # RACI legend
    if is_raci:
        # RACI codes legend
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=4, after=2)
        _add_run(p, "R ", bold=True, size=7, color=RGBColor(0x2E, 0x86, 0xC1))
        _add_run(p, "Responsável  ", size=7, color=COLOR_GRAY)
        _add_run(p, "A ", bold=True, size=7, color=RGBColor(0x27, 0xAE, 0x60))
        _add_run(p, "Aprovador  ", size=7, color=COLOR_GRAY)
        _add_run(p, "C ", bold=True, size=7, color=RGBColor(0xF3, 0x9C, 0x12))
        _add_run(p, "Consultado  ", size=7, color=COLOR_GRAY)
        _add_run(p, "I ", bold=True, size=7, color=RGBColor(0x7F, 0x8C, 0x8D))
        _add_run(p, "Informado", size=7, color=COLOR_GRAY)

        # Person abbreviations legend
        if raci_name_legend:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=2)
            _add_run(p, "Pessoas: ", bold=True, size=7, color=COLOR_GRAY)
            items = [f"{k} = {v}" for k, v in raci_name_legend.items()]
            _add_run(p, " · ".join(items), size=7, color=COLOR_GRAY)

        # Area abbreviations legend
        if raci_area_legend:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=8)
            _add_run(p, "Áreas: ", bold=True, size=7, color=COLOR_GRAY)
            items = [f"{k} = {v}" for k, v in raci_area_legend.items()]
            _add_run(p, " · ".join(items), size=7, color=COLOR_GRAY)

    # Add spacing after table
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Standalone test
# ---------------------------------------------------------------------------
if __name__ == '__main__':
    test_input = Path('/tmp/ata_vallen_abril_2026.md')
    test_output = Path('/tmp/ata_vallen_abril_2026.docx')

    if not test_input.exists():
        print(f"Test input not found: {test_input}")
        sys.exit(1)

    md_content = test_input.read_text(encoding='utf-8')

    # Extract empresa and data from the markdown
    empresa = "VALLEN CLINIC"
    data = "08 de Abril de 2026"

    result = generate_ata_docx(md_content, empresa, data, str(test_output))
    print(f"DOCX generated: {result}")
    print(f"File size: {test_output.stat().st_size:,} bytes")
