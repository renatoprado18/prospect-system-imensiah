#!/usr/bin/env python3
"""
verifica_perfil.py — as superfícies do Renato dizem a mesma coisa?

Compara CV (Drive), peça de posicionamento, site e LinkedIn contra
scripts/fatos_canonicos.json. NÃO corrige nada: aponta divergência.

Nasceu em 12/08/2026, depois que "Osborne, 400 anos" (é 1772) sobreviveu em
quatro documentos. O problema não era falta de cuidado: era não existir um
lugar onde o fato estivesse escrito uma vez.

  python3 scripts/verifica_perfil.py            # todas as superfícies
  python3 scripts/verifica_perfil.py --quiet    # 1 linha (para /cos e /dev)
  python3 scripts/verifica_perfil.py cv peca    # só as indicadas

Saída 0 = tudo alinhado. 1 = há divergência.

⚠️ COBERTURA PARCIAL, E ELA É DITA NO RELATÓRIO. O LinkedIn bloqueia leitura
automática (HTTP 999), então depende de um snapshot manual. Superfície sem
leitura aparece como NÃO VERIFICADA — nunca como "ok". Régua que confunde
"não olhei" com "está certo" certifica o que nunca viu.
"""
import os
import re
import sys
import json
import unicodedata

AQUI = os.path.dirname(os.path.abspath(__file__))
FATOS = os.path.join(AQUI, "fatos_canonicos.json")
RAIZ = os.path.dirname(AQUI)

C = {"erro": "\033[31m", "ok": "\033[32m", "avi": "\033[33m", "dim": "\033[90m",
     "b": "\033[1m", "z": "\033[0m"}
if not sys.stdout.isatty():
    C = {k: "" for k in C}


def norm(t):
    """minúsculas, sem acento, espaços colapsados — para comparar texto de fontes diferentes"""
    t = unicodedata.normalize("NFKD", t.lower())
    t = "".join(c for c in t if not unicodedata.combining(c))
    return re.sub(r"\s+", " ", t)


def env(chave):
    caminho = os.path.join(RAIZ, ".env")
    if os.path.exists(caminho):
        for linha in open(caminho, encoding="utf-8"):
            linha = linha.strip()
            if linha.startswith(chave + "="):
                return linha.split("=", 1)[1].strip().strip('"').strip("'")
    return os.environ.get(chave)


# ─── leitura de cada tipo de superfície ────────────────────────────────

def le_arquivo(cfg):
    p = os.path.expanduser(cfg["caminho"])
    if not os.path.exists(p):
        return None, f"arquivo não existe: {p}"
    s = open(p, encoding="utf-8", errors="replace").read()
    if p.endswith((".html", ".htm")):
        s = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", s, flags=re.S)
        s = re.sub(r"<[^>]+>", " ", s)
    return s, None


def le_url(cfg):
    try:
        import httpx
    except ImportError:
        return None, "httpx não instalado"
    partes = []
    for u in cfg["urls"]:
        try:
            r = httpx.get(u, timeout=25, follow_redirects=True)
            if r.status_code != 200:
                return None, f"{u} respondeu {r.status_code}"
            partes.append(re.sub(r"<[^>]+>", " ", r.text))
        except Exception as e:
            return None, f"{u}: {type(e).__name__}"
    return "\n".join(partes), None


def le_drive_docx(cfg):
    """token do INTEL (google_accounts) + Drive API. O MCP do Drive não serve: não lê docx cru."""
    try:
        import httpx, psycopg2, psycopg2.extras
        from docx import Document
    except ImportError as e:
        return None, f"falta dependência: {e.name}"
    try:
        conn = psycopg2.connect(env("DATABASE_URL_UNPOOLED") or env("DATABASE_URL"))
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute("SELECT access_token FROM google_accounts WHERE tipo='professional' LIMIT 1")
        row = cur.fetchone()
        conn.close()
        if not row:
            return None, "sem conta google 'professional' em google_accounts"
        r = httpx.get(f"https://www.googleapis.com/drive/v3/files/{cfg['file_id']}?alt=media",
                      headers={"Authorization": f"Bearer {row['access_token']}"},
                      timeout=60, follow_redirects=True)
        if r.status_code == 401:
            return None, "token do Google expirado (rode qualquer rota do INTEL para renovar)"
        r.raise_for_status()
    except Exception as e:
        return None, f"{type(e).__name__}: {str(e)[:70]}"

    import io as _io
    doc = Document(_io.BytesIO(r.content))
    partes = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for linha in t.rows:
            partes.extend(c.text for c in linha.cells)
    return "\n".join(partes), None


LEITORES = {"arquivo": le_arquivo, "url": le_url, "drive_docx": le_drive_docx,
            "snapshot": le_arquivo}


# ─── verificação ───────────────────────────────────────────────────────

def main():
    d = json.load(open(FATOS, encoding="utf-8"))
    fatos, superficies = d["fatos"], d["superficies"]

    quiet = "--quiet" in sys.argv
    pedidas = [a for a in sys.argv[1:] if not a.startswith("--")]
    alvos = {k: v for k, v in superficies.items() if not pedidas or k in pedidas}

    textos, crus, naolidas = {}, {}, {}
    for chave, cfg in alvos.items():
        txt, err = LEITORES[cfg["tipo"]](cfg)
        if txt is None:
            naolidas[chave] = err
        else:
            textos[chave] = norm(txt)
            crus[chave] = re.sub(r"\s+", " ", txt)

    achados = []
    for f in fatos:
        for padrao in f["errado"]:
            # Quando o padrao errado so difere do correto por CAIXA ou ACENTO
            # (ex.: "Best For The World" x "Best for the World"), comparar
            # normalizado acusaria o texto certo. Nesses casos, comparacao crua.
            so_caixa = norm(padrao) == norm(f["correto"]) or norm(f["correto"]) in norm(padrao)
            for chave in textos:
                alvo = crus[chave] if so_caixa else textos[chave]
                agulha = padrao if so_caixa else norm(padrao)
                if agulha in alvo:
                    achados.append((chave, f, padrao))

    if quiet:
        n, s, nv = len(achados), len(textos), len(naolidas)
        if achados:
            onde = ", ".join(sorted({a[0] for a in achados}))
            print(f"{C['erro']}🔴 perfil: {n} divergência(s) em {onde}{C['z']}"
                  f"  ·  {s} superfície(s) lida(s), {nv} não verificada(s)")
            return 1
        print(f"{C['ok']}🟢 perfil: {s} superfície(s) alinhada(s){C['z']}"
              + (f"  ·  {C['avi']}{nv} NÃO verificada(s){C['z']}" if nv else ""))
        return 0

    print(f"\n{C['b']}╔═ AS SUPERFÍCIES DIZEM A MESMA COISA? ═╗{C['z']}")
    print(f"{C['dim']}  {len(fatos)} fatos canônicos · fonte: scripts/fatos_canonicos.json"
          f" (atualizado em {d['atualizado_em']}){C['z']}\n")

    for chave, cfg in alvos.items():
        nome = cfg["nome"]
        if chave in naolidas:
            print(f"  {C['avi']}⊘ {nome}: NÃO VERIFICADA{C['z']} — {naolidas[chave]}")
            if cfg.get("como_atualizar"):
                print(f"    {C['dim']}{cfg['como_atualizar']}{C['z']}")
            continue
        meus = [a for a in achados if a[0] == chave]
        if not meus:
            print(f"  {C['ok']}✓ {nome}{C['z']} {C['dim']}— nenhum fato divergente{C['z']}")
        else:
            print(f"  {C['erro']}✗ {nome}: {len(meus)} divergência(s){C['z']}")
            for _, f, padrao in meus:
                print(f"      {C['erro']}achou{C['z']} “{padrao}”")
                print(f"      {C['ok']}correto{C['z']} “{f['correto']}”  {C['dim']}[{f['tema']}]{C['z']}")
                print(f"      {C['dim']}{f['nota']}{C['z']}\n")

    print()
    if naolidas:
        print(f"  {C['avi']}⚠️  {len(naolidas)} de {len(alvos)} superfícies não foram lidas."
              f" Ausência de erro nelas NÃO é prova de acerto.{C['z']}")
    if achados:
        print(f"  {C['erro']}{C['b']}→ {len(achados)} divergência(s). Corrija a superfície,"
              f" ou o fato, se quem mudou foi a verdade.{C['z']}\n")
        return 1
    print(f"  {C['ok']}{C['b']}→ tudo alinhado nas superfícies lidas.{C['z']}\n")
    return 0


if __name__ == "__main__":
    sys.exit(main())
